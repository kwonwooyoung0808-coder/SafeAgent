from __future__ import annotations

from dataclasses import dataclass, field

from docx.document import Document as DocxDocumentObject
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

from src.engines.korean_regulation_parser import KoreanRegulationParser
from src.engines.text_sanitizer import TextSanitizer


@dataclass
class DocxParseResult:
    raw_text:           str
    raw_tables:         list[list[list[str]]]
    doc_structure:      dict
    warnings:           list[str] = field(default_factory=list)
    injection_detected: bool      = False


class DocxEngine:
    """
    python-docx 기반 .docx 파싱 엔진.

    보안:
    - 흰색 폰트(FFFFFF) / vanish 속성의 숨겨진 텍스트 필터링
    - TextSanitizer로 프롬프트 인젝션 패턴 이스케이프 처리

    반환 텍스트는 원본이 아닌 이스케이프된 sanitized_text.
    """

    def parse(self, file_path: str) -> DocxParseResult:
        from docx import Document as DocxDocument

        doc = DocxDocument(file_path)
        raw_text_parts: list[str] = []
        raw_tables: list[list[list[str]]] = []
        blocks: list[dict] = []
        doc_structure: dict = {"headings": [], "tables": [], "blocks": blocks}

        for item in self._iter_document_blocks(doc):
            if isinstance(item, Paragraph):
                if self._is_hidden(item):
                    continue
                text = item.text.strip()
                if not text:
                    continue
                block = {"type": "paragraph", "text": text, "style": item.style.name}
                blocks.append(block)
                if "Heading" in item.style.name:
                    level = int(item.style.name[-1]) if item.style.name[-1].isdigit() else 1
                    doc_structure["headings"].append({"level": level, "text": text})
                raw_text_parts.append(text)
            elif isinstance(item, Table):
                table_rows = self._table_to_rows(item)
                table_text = self._table_to_text(table_rows)
                if not table_text:
                    continue
                raw_tables.append(table_rows)
                doc_structure["tables"].append(table_rows)
                blocks.append({"type": "table", "rows": table_rows, "text": table_text})
                raw_text_parts.append(table_text)

        raw_text = "\n".join(raw_text_parts)
        korean_regulation = KoreanRegulationParser().parse(blocks)
        doc_structure["korean_regulation"] = korean_regulation
        for article in korean_regulation.get("articles", []):
            doc_structure["headings"].append(
                {"level": 2, "text": article["heading"], "source": "korean_article"}
            )

        sanitizer = TextSanitizer()
        sanitized_text, injections = sanitizer.sanitize(raw_text)

        warnings: list[str] = list(injections)
        injection_detected = bool(injections)

        return DocxParseResult(
            raw_text=sanitized_text,
            raw_tables=raw_tables,
            doc_structure=doc_structure,
            warnings=warnings,
            injection_detected=injection_detected,
        )

    @staticmethod
    def _iter_document_blocks(doc: DocxDocumentObject):
        if hasattr(doc, "iter_inner_content"):
            yield from doc.iter_inner_content()
            return
        yield from doc.paragraphs
        yield from doc.tables

    def _table_to_rows(self, table: Table) -> list[list[str]]:
        return [
            [self._cell_to_text(cell) for cell in row.cells]
            for row in table.rows
        ]

    def _cell_to_text(self, cell: _Cell) -> str:
        parts: list[str] = []
        for paragraph in cell.paragraphs:
            if not self._is_hidden(paragraph) and paragraph.text.strip():
                parts.append(paragraph.text.strip())
        for nested in cell.tables:
            nested_text = self._table_to_text(self._table_to_rows(nested))
            if nested_text:
                parts.append(nested_text)
        return "\n".join(parts).strip()

    @staticmethod
    def _table_to_text(rows: list[list[str]]) -> str:
        return "\n".join(
            " | ".join(cell for cell in row if cell)
            for row in rows
            if any(cell for cell in row)
        ).strip()

    @staticmethod
    def _is_hidden(para) -> bool:
        """흰색 폰트(FFFFFF) 또는 OOXML vanish 속성의 숨겨진 텍스트 탐지."""
        NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        for run in para.runs:
            try:
                if run.font.color and run.font.color.type is not None:
                    if str(run.font.color.rgb).upper() in ("FFFFFF", "FFFFFE"):
                        return True
            except Exception:
                pass
            if run._element.find(f".//{{{NS}}}vanish") is not None:
                return True
        return False
