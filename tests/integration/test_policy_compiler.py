"""Feature 3 (POST /v1/policy-compiler/compile) 회귀 테스트.

PRD 5.3.6 수용 기준 검증:
- 변환 실패 시 원본 파일과 오류 로그가 보존된다.
- 변환된 YAML 은 정의된 스키마 검증을 통과한다.
"""
from __future__ import annotations

import io
import uuid
from datetime import date
from pathlib import Path

import pytest
import yaml
from docx import Document


# ══════════════════════════════════════════════════════════════
# F3-1, F3-2 회귀 테스트 — .docx 입력 없이 라우터 검증만
# (실제 .docx 파싱은 python-docx 의존이라 별도 e2e 에서 검증)
# ══════════════════════════════════════════════════════════════


def test_convert_rejects_non_docx_extension(client):
    """라우터 입력 검증: .docx 외 확장자는 400."""
    fake_file = io.BytesIO(b"not a real docx")
    r = client.post(
        "/v1/policy-compiler/compile",
        files={"file": ("policy.txt", fake_file, "text/plain")},
        data={"policy_id": "TEST_POLICY_NON_DOCX", "policy_name": "Test Policy"},
    )
    assert r.status_code == 400, r.text


def test_convert_requires_policy_id(client):
    fake_file = io.BytesIO(b"not a real docx")
    r = client.post(
        "/v1/policy-compiler/compile",
        files={"file": ("policy.docx", fake_file, "application/octet-stream")},
        data={"policy_name": "Missing ID"},
    )
    assert r.status_code == 422, r.text


def test_convert_rejects_invalid_policy_id(client):
    fake_file = io.BytesIO(b"not a real docx")
    r = client.post(
        "/v1/policy-compiler/compile",
        files={"file": ("policy.docx", fake_file, "application/octet-stream")},
        data={"policy_id": "../bad", "policy_name": "Bad ID"},
    )
    assert r.status_code == 422, r.text


def test_convert_rejects_oversized_file(client):
    """라우터 입력 검증: 10MB 초과 파일은 413."""
    big_content = b"a" * (10 * 1024 * 1024 + 1)
    r = client.post(
        "/v1/policy-compiler/compile",
        files={"file": ("big.docx", io.BytesIO(big_content), "application/octet-stream")},
        data={"policy_id": "TEST_POLICY_BIG", "policy_name": "Big"},
    )
    assert r.status_code == 413, r.text


def test_convert_invalid_docx_preserves_original(client, tmp_path, monkeypatch):
    """
    F3-1: 유효하지 않은 .docx 입력 → docx_extractor_node 실패 → status=FAILED.
    원본이 {policy_dir}/failed_uploads/ 에 보존되어야 함.
    F3-2: PolicyConversionLogModel 에 실패 로그가 생성되어야 함.
    """
    # 잘못된 .docx 콘텐츠 (DocxEngine.parse 가 raise 하도록 유도)
    fake_docx = b"PK\x03\x04 not a valid docx zip"

    # policy_dir 가 신규 생성되더라도 안전하게 보존되는지만 확인
    r = client.post(
        "/v1/policy-compiler/compile",
        files={"file": ("malformed.docx", io.BytesIO(fake_docx), "application/octet-stream")},
        data={"policy_id": "TEST_POLICY_MALFORMED", "policy_name": "Malformed"},
    )
    assert r.status_code == 200, r.text
    body = r.json()
    assert body["status"] == "FAILED", body
    # F3-2: failed-{...} 형태의 추적 ID 가 응답에 포함되어야 함
    assert body["policy_id"] is not None, body
    assert body["policy_id"].startswith("failed-"), body
    # F3-1: 원본 보존 경로 안내가 warnings 에 있어야 함
    warnings_text = "\n".join(body["warnings"])
    assert "원본 보존" in warnings_text, body["warnings"]


def test_convert_invalid_docx_creates_conversion_log(client):
    """F3-2: 실패 케이스도 PolicyConversionLogModel 에 기록되어야 함."""
    from src.database.connection import SessionLocal
    from src.database.models import PolicyConversionLogModel

    fake_docx = b"PK\x03\x04 broken"
    r = client.post(
        "/v1/policy-compiler/compile",
        files={"file": ("broken.docx", io.BytesIO(fake_docx), "application/octet-stream")},
        data={"policy_id": "TEST_POLICY_BROKEN", "policy_name": "BrokenPolicy"},
    )
    assert r.status_code == 200
    assert r.json()["status"] == "FAILED"

    # DB 직접 조회로 로그 존재 확인
    session = SessionLocal()
    try:
        rows = session.query(PolicyConversionLogModel).filter(
            PolicyConversionLogModel.original_filename == "broken.docx",
            PolicyConversionLogModel.conversion_status == "FAILED",
        ).all()
        assert len(rows) >= 1, "실패 케이스 PolicyConversionLogModel 누락"
    finally:
        session.close()


def test_convert_uses_required_policy_id_and_server_effective_date(client):
    from src.database.connection import SessionLocal
    from src.database.models import PolicyModel

    doc = Document()
    doc.add_paragraph("제1조(목적) 회사의 영업비밀을 보호하여야 한다.")
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    policy_id = f"TEST_EFFECTIVE_{uuid.uuid4().hex[:8].upper()}"
    r = client.post(
        "/v1/policy-compiler/compile",
        files={
            "file": (
                "valid.docx",
                buf,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        data={"policy_id": policy_id, "policy_name": "Effective Date Policy"},
    )
    assert r.status_code == 200, r.text
    body = r.json()
    assert body["policy_id"] == policy_id
    assert body["yaml_path"]

    session = SessionLocal()
    try:
        row = session.query(PolicyModel).filter(PolicyModel.id == policy_id).first()
        assert row is not None
        assert row.effective_date == date.today()
    finally:
        session.close()
        Path(body["yaml_path"]).unlink(missing_ok=True)


def test_compile_creates_initial_inactive_version(client):
    doc = Document()
    doc.add_paragraph("제1조(목적) 회사의 영업비밀을 보호하여야 한다.")
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    policy_id = f"TEST_VERSION_{uuid.uuid4().hex[:8].upper()}"
    r = client.post(
        "/v1/policy-compiler/compile",
        files={
            "file": (
                "versioned.docx",
                buf,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        data={"policy_id": policy_id, "policy_name": "Versioned Policy"},
    )
    assert r.status_code == 200, r.text
    body = r.json()

    versions = client.get(f"/v1/policy-compiler/{policy_id}/versions").json()
    assert versions["total"] == 1
    assert versions["items"][0]["version"] == "1.0"
    assert versions["items"][0]["is_current"] is False
    Path(body["yaml_path"]).unlink(missing_ok=True)


def test_activation_blocks_unreviewed_draft(client):
    doc = Document()
    doc.add_paragraph("제1조(목적) 회사의 영업비밀을 보호하여야 한다.")
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    policy_id = f"TEST_REVIEW_{uuid.uuid4().hex[:8].upper()}"
    r = client.post(
        "/v1/policy-compiler/compile",
        files={
            "file": (
                "review.docx",
                buf,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        data={"policy_id": policy_id, "policy_name": "Review Policy"},
    )
    assert r.status_code == 200, r.text
    yaml_path = Path(r.json()["yaml_path"])

    activate = client.put(f"/v1/policy-compiler/{policy_id}/activate")
    assert activate.status_code == 422

    draft = client.get(f"/v1/policy-compiler/{policy_id}/draft")
    assert draft.status_code == 200
    body = draft.json()
    assert body["needs_review"] is True

    reviewed_doc = yaml.safe_load(body["raw_yaml"])
    reviewed_doc["judge"]["criteria"] = "관리자가 검토를 완료한 정책 기준입니다."
    reviewed_yaml = yaml.safe_dump(
        reviewed_doc,
        allow_unicode=True,
        default_flow_style=False,
        sort_keys=False,
    )
    update = client.put(
        f"/v1/policy-compiler/{policy_id}/draft",
        json={"raw_yaml": reviewed_yaml},
    )
    assert update.status_code == 200, update.text
    assert update.json()["needs_review"] is False
    assert update.json()["version"] == "1.1"

    activate = client.put(f"/v1/policy-compiler/{policy_id}/activate")
    assert activate.status_code == 200, activate.text

    versions = client.get(f"/v1/policy-compiler/{policy_id}/versions").json()
    currents = [item for item in versions["items"] if item["is_current"]]
    assert len(currents) == 1
    assert currents[0]["version"] == "1.1"
    yaml_path.unlink(missing_ok=True)


def _create_reviewed_policy(client, policy_id: str) -> Path:
    doc = Document()
    doc.add_paragraph("제1조(목적) 회사의 영업비밀을 보호하여야 한다.")
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    r = client.post(
        "/v1/policy-compiler/compile",
        files={
            "file": (
                f"{policy_id}.docx",
                buf,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        data={"policy_id": policy_id, "policy_name": policy_id},
    )
    assert r.status_code == 200, r.text
    yaml_path = Path(r.json()["yaml_path"])
    draft = client.get(f"/v1/policy-compiler/{policy_id}/draft").json()
    doc_yaml = yaml.safe_load(draft["raw_yaml"])
    doc_yaml["judge"]["criteria"] = "관리자가 검토를 완료한 정책 기준입니다."
    reviewed_yaml = yaml.safe_dump(
        doc_yaml,
        allow_unicode=True,
        default_flow_style=False,
        sort_keys=False,
    )
    update = client.put(
        f"/v1/policy-compiler/{policy_id}/draft",
        json={"raw_yaml": reviewed_yaml},
    )
    assert update.status_code == 200, update.text
    return yaml_path


def test_deactivate_policy_preserves_yaml_and_clears_current_version(client):
    policy_id = f"TEST_DEACT_{uuid.uuid4().hex[:8].upper()}"
    yaml_path = _create_reviewed_policy(client, policy_id)

    activate = client.put(f"/v1/policy-compiler/{policy_id}/activate")
    assert activate.status_code == 200, activate.text
    deactivate = client.put(f"/v1/policy-compiler/{policy_id}/deactivate")
    assert deactivate.status_code == 200, deactivate.text
    assert deactivate.json()["is_active"] is False
    assert yaml_path.exists()

    versions = client.get(f"/v1/policy-compiler/{policy_id}/versions").json()
    assert all(item["is_current"] is False for item in versions["items"])
    yaml_path.unlink(missing_ok=True)


def test_delete_unused_policy_removes_db_rows_and_yaml(client):
    policy_id = f"TEST_DELETE_{uuid.uuid4().hex[:8].upper()}"
    yaml_path = _create_reviewed_policy(client, policy_id)

    delete = client.delete(f"/v1/policy-compiler/{policy_id}")
    assert delete.status_code == 200, delete.text
    assert delete.json()["deleted"] is True
    assert not yaml_path.exists()
    assert client.get(f"/v1/policy-compiler/{policy_id}").status_code == 404
    assert client.get(f"/v1/policy-compiler/{policy_id}/versions").status_code == 404


def test_delete_blocks_policy_used_by_group(client):
    policy_id = f"TEST_GROUP_{uuid.uuid4().hex[:8].upper()}"
    yaml_path = _create_reviewed_policy(client, policy_id)
    activate = client.put(f"/v1/policy-compiler/{policy_id}/activate")
    assert activate.status_code == 200, activate.text
    group_id = f"GRP_{uuid.uuid4().hex[:8].upper()}"
    group = client.post(
        "/v1/policy-groups",
        json={"id": group_id, "name": group_id, "policy_ids": [policy_id]},
    )
    assert group.status_code == 201, group.text
    deactivate = client.put(f"/v1/policy-compiler/{policy_id}/deactivate")
    assert deactivate.status_code == 200, deactivate.text

    delete = client.delete(f"/v1/policy-compiler/{policy_id}")
    assert delete.status_code == 409
    assert "policy_groups" in delete.json()["detail"]["blockers"]
    assert yaml_path.exists()
    yaml_path.unlink(missing_ok=True)
